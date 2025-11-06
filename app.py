# ==========================================================
# 📚 Multi-Language Audiobook & Text Processor (Final Version)
# ==========================================================

import streamlit as st
import docx
import fitz  # PyMuPDF for PDF reading
import tempfile
import os
from gtts import gTTS
import base64

# ===================== Optional Dependencies =====================
try:
    import nltk
    nltk.download('punkt')
    from nltk.tokenize import word_tokenize
    NLTK_AVAILABLE = True
except:
    NLTK_AVAILABLE = False

# ===================== Utility Functions =====================

def extract_text(file):
    """Extract text from TXT, DOCX, or PDF file."""
    name = file.name.lower()

    if name.endswith(".txt"):
        return file.read().decode("utf-8", errors="ignore")

    elif name.endswith(".docx"):
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])

    elif name.endswith(".pdf"):
        # Save PDF temporarily to read using fitz
        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
            tmp.write(file.read())
            tmp_path = tmp.name

        text = ""
        pdf = fitz.open(tmp_path)
        for page in pdf:
            text += page.get_text("text") or ""
        pdf.close()
        os.remove(tmp_path)
        return text

    else:
        raise ValueError("❌ Unsupported file type. Please upload txt, docx, or pdf.")


def tokenize_text(text):
    """Tokenize text into words using NLTK or fallback split."""
    if NLTK_AVAILABLE:
        return word_tokenize(text)
    else:
        return text.split()


def chunk_text(text, chunk_size=3000, overlap=200):
    """Split text into overlapping chunks for TTS."""
    words = text.split()
    chunks = []
    i = 0
    while i < len(words):
        chunk = words[i:i + chunk_size]
        chunks.append(" ".join(chunk))
        i += chunk_size - overlap
    return chunks


def generate_tts(text, lang="en"):
    """Convert text into MP3 using gTTS and return the file path."""
    tts = gTTS(text, lang=lang)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".mp3")
    tts.save(tmp.name)
    return tmp.name

# ===================== Streamlit UI =====================

st.set_page_config(
    page_title="📚 Multi-Language Audiobook & Text Processor",
    layout="centered"
)

st.title("📚 Multi-Language Audiobook & Text Processor")
st.caption("Upload your document → Tokenize → Chunk → Generate Audio 🎧")

# ---------- Upload Section ----------
uploaded_file = st.file_uploader("📤 Upload a text / pdf / docx file", type=["txt", "pdf", "docx"])

# ---------- Parameters ----------
st.markdown("### ⚙️ Settings")
chunk_size = st.number_input("Chunk size (words)", min_value=500, max_value=8000, value=2000, step=500)
overlap = st.number_input("Chunk overlap (words)", min_value=0, max_value=500, value=200, step=50)
tts_lang = st.selectbox("🎧 Select Audio Language", ["en", "hi", "fr", "es"])

# ===================== MAIN LOGIC =====================

if uploaded_file:
    text = extract_text(uploaded_file)

    # --- Preview Text ---
    st.subheader("📜 Original Text Preview")
    st.text_area("", text[:2000], height=200)

    # --- Tokenization ---
    st.subheader("🔠 Tokenization")
    tokens = tokenize_text(text)
    st.write(f"**Total Tokens:** {len(tokens)}")
    st.write(tokens[:50])  # Show first few tokens

    # --- Chunking ---
    st.subheader("✂️ Text Chunking")
    chunks = chunk_text(text, chunk_size, overlap)
    st.write(f"**Total Chunks:** {len(chunks)}")
    st.write(chunks[0][:500] + "..." if chunks else "⚠️ No chunks found.")

    # --- Audio Generation ---
    if st.button("🎧 Generate Audiobook"):
        with st.spinner("🎤 Generating audio, please wait..."):
            for i, c in enumerate(chunks, 1):
                audio_path = generate_tts(c, lang=tts_lang)
                st.audio(audio_path)  # Play audio in the app

                # Provide download link
                with open(audio_path, "rb") as f:
                    b64 = base64.b64encode(f.read()).decode()
                    href = f'<a href="data:audio/mp3;base64,{b64}" download="chunk_{i}.mp3">📥 Download Chunk {i}</a>'
                    st.markdown(href, unsafe_allow_html=True)

        st.success("✅ Audiobook generation completed successfully!")

else:
    st.info("📂 Please upload a file to begin.")







# ============================================================
# 🌍 + 🔊 MULTI-LANGUAGE TRANSLATOR & AUDIOBOOK GENERATOR (v3.1 FIXED)
# ============================================================

import streamlit as st
from deep_translator import GoogleTranslator
from gtts import gTTS
from datetime import datetime
import tempfile
import base64
import os

# ✅ Page setup
st.set_page_config(page_title="Translator + Audiobook Generator", page_icon="🌍")
st.title("🌍 + 🔊 Multi-Language Translator & Audiobook Generator")

# ✅ Supported languages
lang_dict = {
    "Hindi": "hi", "Bengali": "bn", "Tamil": "ta", "Telugu": "te",
    "Gujarati": "gu", "Malayalam": "ml", "Kannada": "kn", "Marathi": "mr",
    "Urdu": "ur", "English": "en",
    "Chhattisgarhi (Experimental)": "hne"
}

# ✅ Initialize session state
if "translated_text" not in st.session_state:
    st.session_state.translated_text = ""

# ============================================================
# 🌍 TEXT TRANSLATION SECTION
# ============================================================
st.header("🌐 Step 1: Translate Your Text")

# ✅ Text input
text_input = st.text_area("✍️ Enter text to translate", 
                          "Type or paste English text here...", height=150)

# ✅ Language selection
target_lang = st.selectbox("🌍 Choose Target Language", list(lang_dict.keys()), index=0)

# ✅ Translate button
if st.button("🔁 Translate Text"):
    if not text_input.strip():
        st.warning("⚠️ Please enter some text to translate.")
    else:
        try:
            translated_text = GoogleTranslator(source='auto', target=lang_dict[target_lang]).translate(text_input)
            st.session_state.translated_text = translated_text  # ✅ Store in session
            st.success(f"✅ Translated Text ({target_lang}):")
            st.markdown(f"```\n{translated_text}\n```")
        except Exception as e:
            st.error(f"❌ Translation Error: {e}")

# ============================================================
# 🔊 AUDIOBOOK GENERATION SECTION
# ============================================================
st.header("🎧 Step 2: Generate Audio from Translated Text")

if not st.session_state.translated_text:
    st.info("ℹ️ Please translate your text first to enable audio generation.")

if st.button("🎙️ Generate Audio"):
    if not st.session_state.translated_text.strip():
        st.warning("⚠️ Please translate some text first.")
    else:
        try:
            translated_text = st.session_state.translated_text
            st.write(f"🎧 Generating audio in **{target_lang}** ...")
            tts = gTTS(translated_text, lang=lang_dict[target_lang])

            # Save with timestamp
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            file_name = f"translated_audio_{target_lang}_{timestamp}.mp3"
            audio_path = os.path.join(tempfile.gettempdir(), file_name)
            tts.save(audio_path)

            # Play audio
            st.audio(audio_path, format='audio/mp3', start_time=0)

            # Download link
            with open(audio_path, "rb") as f:
                audio_bytes = f.read()
            b64 = base64.b64encode(audio_bytes).decode()
            href = f'<a href="data:audio/mp3;base64,{b64}" download="{file_name}" style="text-decoration:none;"><button style="background-color:#4CAF50;color:white;padding:10px 15px;border:none;border-radius:8px;cursor:pointer;">⬇️ Download {target_lang} Audio</button></a>'
            st.markdown(href, unsafe_allow_html=True)

        except Exception as e:
            st.error(f"❌ Error generating audio: {e}")



















 # ============================================================
# 🌍 Universal Translator + Audiobook Generator (Streamlit)
# ============================================================

import streamlit as st
import os
import tempfile
from gtts import gTTS
from deep_translator import GoogleTranslator
import PyPDF2
import docx
from odf.opendocument import load
from odf import text, teletype
from datetime import datetime

# -------------------------------
# 🌐 Supported Languages
# -------------------------------
languages = {
    "en": "English",
    "hi": "Hindi",
    "bn": "Bengali",
    "ta": "Tamil",
    "te": "Telugu",
    "mr": "Marathi",
    "gu": "Gujarati",
    "kn": "Kannada",
    "ml": "Malayalam",
    "pa": "Punjabi",
    "or": "Odia",
    "ur": "Urdu",
    "ne": "Nepali",
    "bho": "Bhojpuri (via Hindi)",
    "cg": "Chhattisgarhi (via Hindi)"
}

# -------------------------------
# Streamlit UI
# -------------------------------
st.set_page_config(page_title="Universal Translator + Audiobook", page_icon="🎧")
st.title("🌍 Universal Translator + Audiobook Generator")
st.markdown("Upload a file (TXT, PDF, DOCX, ODT), translate it to your chosen language, and generate audio!")

uploaded_file = st.file_uploader("📂 Upload File", type=["txt", "pdf", "docx", "odt"])
target_lang = st.selectbox("🎯 Select Target Language", list(languages.values()))

# -------------------------------
# Extract Text
# -------------------------------
def extract_text(file):
    ext = os.path.splitext(file.name)[1].lower()
    text_content = ""

    if ext == ".txt":
        text_content = str(file.read(), "utf-8")
    elif ext == ".pdf":
        reader = PyPDF2.PdfReader(file)
        for page in reader.pages:
            text_content += page.extract_text() or ""
    elif ext == ".docx":
        doc = docx.Document(file)
        for para in doc.paragraphs:
            text_content += para.text + "\n"
    elif ext == ".odt":
        doc = load(file)
        allparas = doc.getElementsByType(text.P)
        for para in allparas:
            text_content += teletype.extractText(para) + "\n"
    else:
        st.error("❌ Unsupported file type!")
    return text_content.strip()

# -------------------------------
# Map language name → code
# -------------------------------
def get_lang_code(lang_name):
    lang_name = lang_name.lower()
    # Custom mappings
    if lang_name in ["bhojpuri", "bhojpuri (via hindi)", "bho"]:
        return "hi"
    if lang_name in ["chhattisgarhi", "chhattisgarhi (via hindi)", "cg"]:
        return "hi"
    # normal case
    for code, name in languages.items():
        if name.lower() == lang_name:
            return code
    return "en"  # default English

# -------------------------------
# Process File
# -------------------------------
if uploaded_file:
    try:
        extracted_text = extract_text(uploaded_file)
        if not extracted_text:
            st.warning("⚠️ No text found in file.")
            st.stop()

        st.subheader("📖 Extracted Text")
        st.write(extracted_text[:2000] + ("..." if len(extracted_text) > 2000 else ""))

        target_code = get_lang_code(target_lang)

        # -------------------------------
        # Translate paragraph by paragraph
        # -------------------------------
        st.info("🔄 Translating text...")
        try:
            paragraphs = [p for p in extracted_text.split("\n") if p.strip() != ""]
            translated_paragraphs = []
            for para in paragraphs:
                try:
                    translated_para = GoogleTranslator(source='auto', target=target_code).translate(para)
                except Exception:
                    translated_para = para  # fallback if translation fails
                translated_paragraphs.append(translated_para)
            translated_text = "\n".join(translated_paragraphs)
            st.success("✅ Translation Complete!")
        except Exception as e:
            st.error(f"❌ Translation Failed: {e}")
            translated_text = extracted_text  # fallback: original text

        st.subheader(f"🌐 Translated Text ({target_lang})")
        st.write(translated_text[:2000] + ("..." if len(translated_text) > 2000 else ""))

        # -------------------------------
        # Generate TTS Audio
        # -------------------------------
        st.info("🎧 Generating audio...")
        try:
            tts = gTTS(translated_text, lang=target_code)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            audio_filename = f"translated_audio_{target_lang}_{ts}.mp3"
            audio_path = os.path.join(tempfile.gettempdir(), audio_filename)
            tts.save(audio_path)
            st.success("✅ Audio Generated!")

            # Play Audio
            st.audio(audio_path)

            # Download Audio
            with open(audio_path, "rb") as f:
                st.download_button(
                    label="💾 Download Audio (MP3)",
                    data=f,
                    file_name=audio_filename,
                    mime="audio/mp3"
                )
        except Exception as e:
            st.error(f"❌ TTS Generation Failed: {e}")

        # Download Translated Text
        text_filename = f"translated_text_{target_lang}_{ts}.txt"
        text_path = os.path.join(tempfile.gettempdir(), text_filename)
        with open(text_path, "w", encoding="utf-8") as f:
            f.write(translated_text)
        with open(text_path, "rb") as f:
            st.download_button(
                label="💾 Download Translated Text",
                data=f,
                file_name=text_filename,
                mime="text/plain"
            )

    except Exception as e:
        st.error(f"❌ Error: {e}")







import streamlit as st
from transformers import pipeline
from deep_translator import GoogleTranslator
import torch

# --------- Setup ---------
device = 0 if torch.cuda.is_available() else -1

st.set_page_config(page_title="🌟 Emotion Detector", page_icon="😄")
st.title("😄 Emotion Detection + Multi-language Support")
st.markdown("Enter text in any language, and detect the dominant emotion with explanation.")

# Load model
@st.cache_resource(show_spinner=True)
def load_emotion_model():
    return pipeline(
        "text-classification",
        model="j-hartmann/emotion-english-distilroberta-base",
        return_all_scores=True,
        device=device
    )

emotion_classifier = load_emotion_model()
translator = GoogleTranslator(source='auto', target='en')

# --------- Function to detect emotion ---------
def detect_emotion(text):
    if not text.strip():
        st.warning("⚠️ Empty text!")
        return

    # Translate text to English
    text_en = translator.translate(text)
    st.markdown(f"🌐 **Translated to English:** {text_en}")

    # Split text into chunks of 500 characters
    chunks = [text_en[i:i+500] for i in range(0, len(text_en), 500)]
    combined_scores = {}

    for chunk in chunks:
        result = emotion_classifier(chunk)[0]
        for r in result:
            combined_scores[r['label']] = combined_scores.get(r['label'], 0) + r['score']

    # Average scores
    for k in combined_scores:
        combined_scores[k] /= len(chunks)

    # Sort & get dominant emotion
    sorted_scores = sorted(combined_scores.items(), key=lambda x: x[1], reverse=True)
    label, score = sorted_scores[0]

    mood_map = {
        "joy": "😊 You seem happy and content.",
        "sadness": "😢 You sound sad or reflective.",
        "anger": "😡 You might be feeling angry or frustrated.",
        "fear": "😨 There’s some fear or anxiety in your tone.",
        "love": "❤️ You’re expressing affection or care.",
        "surprise": "😲 There’s surprise or shock in your tone.",
        "disgust": "🤢 There’s dislike or rejection in your words.",
        "neutral": "😐 Your tone feels calm and neutral."
    }

    mood_summary = mood_map.get(label.lower(), "🙂 General emotion detected.")
    
    st.subheader(f"🎯 Dominant Emotion: {label} (confidence: {score:.3f})")
    st.info(mood_summary)

    st.subheader("📊 All emotions ranked:")
    for lbl, scr in sorted_scores:
        st.write(f"- {lbl:<10}: {scr:.3f}")

# --------- User Input ---------
user_text = st.text_area("📝 Enter your text (any language):", height=150)

if st.button("Detect Emotion"):
    detect_emotion(user_text)

      


# ============================================================
# 📝 FOOTER
# ============================================================
st.markdown("---")
st.markdown("💡 *Developed by Aman | Powered by Deep Translator + gTTS*")
